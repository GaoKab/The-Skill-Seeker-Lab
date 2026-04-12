"""
Export NodeRail strategy documents to Word (.docx) files.

Run: python3 export_docs.py

Outputs:
  - noderail_positioning.docx
  - noderail_terms_of_service.docx
  - noderail_pricing.docx
  - noderail_licensing.docx
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)


# ===================================================================
# 1. POSITIONING
# ===================================================================

def export_positioning():
    from noderail.positioning import (
        CATEGORY, TAGLINES, ELEVATOR_PITCH_30_SEC, ELEVATOR_PITCH_2_MIN,
        COMPETITORS, POSITIONING_MATRIX, POSITIONING_STATEMENT,
    )

    doc = Document()
    set_style(doc)

    doc.add_heading('NodeRail™ — Competitor Positioning & Messaging', 0)
    doc.add_paragraph('Strategic reference document. Confidential.', style='Intense Quote')

    # Category
    doc.add_heading('Category Definition', level=1)
    doc.add_heading(CATEGORY['name'], level=2)
    doc.add_paragraph(CATEGORY['definition'])
    doc.add_heading('Why a new category?', level=3)
    doc.add_paragraph(CATEGORY['why_new_category'])

    # Taglines
    doc.add_heading('Tagline Options', level=1)
    for t in TAGLINES:
        doc.add_paragraph(t, style='List Bullet')

    # Elevator pitches
    doc.add_heading('Elevator Pitch — 30 Seconds', level=1)
    doc.add_paragraph(ELEVATOR_PITCH_30_SEC)

    doc.add_heading('Elevator Pitch — 2 Minutes', level=1)
    for para in ELEVATOR_PITCH_2_MIN.split('\n\n'):
        doc.add_paragraph(para)

    # Competitors
    doc.add_heading('Competitor Positioning', level=1)
    for name, data in COMPETITORS.items():
        display_name = name.replace('_', ' ').title()
        doc.add_heading(f'vs. {display_name}', level=2)

        doc.add_heading('One-liner', level=3)
        doc.add_paragraph(data['differentiator'])

        doc.add_heading('What they do well', level=3)
        doc.add_paragraph(data['does_well'])

        doc.add_heading('What they lack', level=3)
        doc.add_paragraph(data['lacks'])

        doc.add_heading('When someone asks', level=3)
        doc.add_paragraph(data['when_asked'])

    # Matrix
    doc.add_heading('Positioning Matrix', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(POSITIONING_MATRIX['header']):
        table.rows[0].cells[i].text = h
    for row_data in POSITIONING_MATRIX['rows']:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    # Positioning statement
    doc.add_heading('Positioning Statement', level=1)
    doc.add_paragraph(POSITIONING_STATEMENT)

    doc.save('noderail_positioning.docx')
    print('  + noderail_positioning.docx')


# ===================================================================
# 2. TERMS OF SERVICE
# ===================================================================

def export_terms():
    from noderail.terms_outline import (
        CONTENT_OWNERSHIP, REVIEWER_LIABILITY, AI_API_USAGE,
        DATA_STORAGE_PRIVACY, INSTITUTIONAL_DATA_AGREEMENTS,
        ACCOUNT_TERMINATION, DOI_PERMANENCE,
        TERMS_SECTIONS, ADDITIONAL_SECTIONS_TODO,
    )

    sections = {
        'CONTENT_OWNERSHIP': CONTENT_OWNERSHIP,
        'REVIEWER_LIABILITY': REVIEWER_LIABILITY,
        'AI_API_USAGE': AI_API_USAGE,
        'DATA_STORAGE_PRIVACY': DATA_STORAGE_PRIVACY,
        'INSTITUTIONAL_DATA_AGREEMENTS': INSTITUTIONAL_DATA_AGREEMENTS,
        'ACCOUNT_TERMINATION': ACCOUNT_TERMINATION,
        'DOI_PERMANENCE': DOI_PERMANENCE,
    }

    doc = Document()
    set_style(doc)

    doc.add_heading('NodeRail™ — Terms of Service (Draft)', 0)
    doc.add_paragraph(
        'Structural outline for legal counsel. Not a final legal document. '
        'Every section should be reviewed by an attorney before publication.',
        style='Intense Quote',
    )

    for info in TERMS_SECTIONS:
        section_data = sections[info['key']]

        doc.add_heading(f'{info["number"]}. {info["title"]}', level=1)

        # Principle
        doc.add_heading('Guiding Principle', level=2)
        doc.add_paragraph(section_data['principle'])

        # Provisions
        doc.add_heading('Provisions', level=2)
        for prov in section_data['provisions']:
            doc.add_heading(prov['title'], level=3)
            doc.add_paragraph(prov['detail'])

    # Additional sections
    doc.add_heading('Additional Sections (To Be Addressed)', level=1)
    for item in ADDITIONAL_SECTIONS_TODO:
        doc.add_paragraph(item, style='List Bullet')

    doc.save('noderail_terms_of_service.docx')
    print('  + noderail_terms_of_service.docx')


# ===================================================================
# 3. PRICING
# ===================================================================

def export_pricing():
    from noderail.pricing import (
        FREE_TIER, PRO_TIER, INSTITUTIONAL_TIER,
        AI_API_TIER, ENTERPRISE_TIER,
        PRICING_SUMMARY, STRATEGIC_NOTES,
    )

    doc = Document()
    set_style(doc)

    doc.add_heading('NodeRail™ — Pricing Tiers', 0)
    doc.add_paragraph('Strategic reference document. Confidential.', style='Intense Quote')

    # Summary table
    doc.add_heading('Pricing Summary', level=1)
    table = doc.add_table(rows=1, cols=len(PRICING_SUMMARY['header']))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(PRICING_SUMMARY['header']):
        table.rows[0].cells[i].text = h
    for row_data in PRICING_SUMMARY['rows']:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    # Tiers
    for tier_data, name in [
        (FREE_TIER, 'Free Tier'),
        (PRO_TIER, 'Pro Tier'),
    ]:
        doc.add_heading(f'{name} — {tier_data["price"]}', level=1)
        doc.add_paragraph(f'Target: {tier_data["target"]}')

        doc.add_heading('Includes', level=2)
        for item in tier_data['includes']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('Limits', level=2)
        for item in tier_data['limits']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('Rationale', level=2)
        doc.add_paragraph(tier_data['rationale'])

    # Institutional sub-tiers
    doc.add_heading('Institutional Tiers', level=1)
    for key, sub in INSTITUTIONAL_TIER['sub_tiers'].items():
        doc.add_heading(f'{sub["name"]} — {sub["price"]}', level=2)
        doc.add_paragraph(f'Target: {sub["target"]}')
        doc.add_paragraph(f'Seats: {sub["seats"]}')

        doc.add_heading('Includes', level=3)
        for item in sub['includes']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('Rationale', level=3)
        doc.add_paragraph(sub['rationale'])

    # AI API
    doc.add_heading(f'AI API Tier — {AI_API_TIER["price"]}', level=1)
    doc.add_paragraph(f'Target: {AI_API_TIER["target"]}')

    doc.add_heading('Rate Card', level=2)
    for tier_key, tier_info in AI_API_TIER['rate_card'].items():
        doc.add_paragraph(
            f'{tier_key.replace("_", " ").title()}: {tier_info["price"]} — '
            f'{tier_info["limits"]}',
            style='List Bullet',
        )

    doc.add_heading('API Features', level=2)
    for item in AI_API_TIER['api_features']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Rationale', level=2)
    doc.add_paragraph(AI_API_TIER['rationale'])

    # Enterprise
    doc.add_heading(f'Enterprise Tier — {ENTERPRISE_TIER["price"]}', level=1)
    doc.add_paragraph(f'Target: {ENTERPRISE_TIER["target"]}')

    doc.add_heading('Includes', level=2)
    for item in ENTERPRISE_TIER['includes']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Rationale', level=2)
    doc.add_paragraph(ENTERPRISE_TIER['rationale'])

    # Strategic notes
    doc.add_heading('Strategic Notes', level=1)
    for key, val in STRATEGIC_NOTES.items():
        if isinstance(val, str):
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            doc.add_paragraph(val)
        elif isinstance(val, list):
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            for item in val:
                doc.add_paragraph(item, style='List Bullet')
        elif isinstance(val, dict):
            doc.add_heading(key.replace('_', ' ').title(), level=2)
            for k2, v2 in val.items():
                doc.add_paragraph(f'{k2}: {v2}', style='List Bullet')

    doc.save('noderail_pricing.docx')
    print('  + noderail_pricing.docx')


# ===================================================================
# 4. LICENSING
# ===================================================================

def export_licensing():
    from noderail.licensing import LICENSING_STRATEGY

    doc = Document()
    set_style(doc)

    doc.add_heading('NodeRail™ — Licensing Strategy', 0)
    doc.add_paragraph('Strategic reference document. Confidential.', style='Intense Quote')

    doc.add_heading('Decision', level=1)
    doc.add_paragraph(LICENSING_STRATEGY['decision'])

    doc.add_heading('Open Components (CC BY 4.0)', level=1)
    for key, data in LICENSING_STRATEGY['open_components'].items():
        doc.add_heading(key.replace('_', ' ').title(), level=2)
        doc.add_paragraph(f'License: {data["license"]}')
        doc.add_paragraph(data['rationale'])

    doc.add_heading('Proprietary Components', level=1)
    for key, data in LICENSING_STRATEGY['proprietary_components'].items():
        doc.add_heading(key.replace('_', ' ').title(), level=2)
        doc.add_paragraph(f'What: {data["what"]}')
        doc.add_paragraph(data['rationale'])

    doc.add_heading('Trademark', level=1)
    doc.add_paragraph(f'Name: {LICENSING_STRATEGY["trademark"]["name"]}')
    doc.add_paragraph(f'Status: {LICENSING_STRATEGY["trademark"]["status"]}')
    doc.add_heading('Action Items', level=2)
    for item in LICENSING_STRATEGY['trademark']['action_items']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Summary', level=1)
    doc.add_paragraph(LICENSING_STRATEGY['summary'])

    doc.save('noderail_licensing.docx')
    print('  + noderail_licensing.docx')


# ===================================================================
# RUN ALL
# ===================================================================

if __name__ == '__main__':
    print('Exporting NodeRail strategy documents to Word...\n')
    export_positioning()
    export_terms()
    export_pricing()
    export_licensing()
    print('\nDone. Files are in the current directory.')
